import docx
from docx import Document
from docx.shared import RGBColor
import re
from tkinter import *
from tkinter import filedialog
from tkinter import ttk
from tkinter.messagebox import showinfo
import os
import xlwings

# DER and TBV are not valid tags
docRelation = {"HRD":("HRS"), "HRS":("PRS"), "PRS":("URS","RISK"), "HTR":("HTP"), "HTP":("HRD", "HRS"), \
               "SDS":("BOLUS","ACE","AID"), "ACE":("PRS"), "BOLUS":("PRS"), "AID":("PRS"), \
               "SVAL":("BOLUS", "ACE", "AID"), "SVATR":("SVAL"), "UT":("UNIT"), "INS": ("UNIT")}      # to be created by the GUI

docFile = {"HRD":"HDS_new_pump.docx", "HRS":"HRS_new_pump.docx", "HTP":"HTP_new_pump.docx", "HTR":"HTR_new_pump.docx", \
           "PRS":"PRS_new_pump.docx", "RISK":"RiskAnalysis_Pump.docx", "SDS":"SDS_New_pump_x04.docx", \
           "ACE":"SRS_ACE_Pump_X01.docx", "BOLUS":"SRS_BolusCalc_Pump_X04.docx", "SRS":"SRS_DosingAlgorithm_X03.docx", \
           "SVAL":"SVaP_new_pump.docx", "SVATR":"SVaTR_new_pump.docx", "UT":"SVeTR_new_pump.docx", "URS":"URS_new_pump.docx"}

docFileList = list(docFile.keys())                  # This is a list of all main tags found in each document
print(docFileList)
parentTagList = list(docRelation.values())

report3 = Document()                #create word document
paragraph = report3.add_paragraph()
report3.save('report3.docx')

uniqueValidTagList = []                             # This is the valid child tag list
for tag in parentTagList:
    if type(tag) is tuple:                          # if a tuple is found, convert to a list and add to the list
        uniqueValidTagList.extend(list(tag))
    else:
        uniqueValidTagList.append(tag)              # if not a tuple simply append to the list
uniqueTagList = (list(set(uniqueValidTagList)))     # set() strips out all redundant tags


#filePath = "C:/Users/steph/OneDrive/Desktop/Docs_Project/"
filePath = "/Users/adrian/Desktop/SampledocsTandem/"

def GetText(filename):                      # Opens the document and places each paragraph into a list
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    fullText = [ele for ele in fullText if ele.strip()]   # Eliminates empty paragraphs
    return fullText

def GetParentTags():                    # Returns only valid parent tags
    for tag in docFileList:             # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + "SRS" + '[:\s]\S*', t)
                    red = paragraph.add_run(y)
                    paragraph.add_run("\n\n")
                    red.bold = True
                    red.font.color.rgb = RGBColor(255, 0, 0)
                    #print(y[0])
                index = index + 1
            # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + re.escape(tag) + '[:\s]\S*', t)
                    red = paragraph.add_run(y)
                    paragraph.add_run("\n\n")
                    red.bold = True
                    red.font.color.rgb = RGBColor(255, 0, 0)
                    #print(y[0])
                index = index + 1
            #print(ind)

def GetChildTags():                     # Returns only valid child tags
    for tag in docFileList:             # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\[.+\]', t)
                    if len(y) != 0:
                        green = paragraph.add_run(y[0])
                        paragraph.add_run("\n\n")
                        green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        green.bold = True
                        #print(y[0])
                index = index + 1
                # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\[.+\]', t)
                    if len(y) != 0:
                        green = paragraph.add_run(y[0])
                        paragraph.add_run("\n\n")
                        green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        green.bold = True
                        #print(y[0])
                index = index + 1
            #print(ind)


def GetOrphanTags():
    for tag in docFileList:  # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            #y = re.findall('[\s\]]\[.+\][\[\s]', t)
            y = re.findall('\[.+\]', t)
            if len(y) != 0:

                #green = paragraph.add_run(y[0])
                #paragraph.add_run("\n\n")
                #green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                #green.bold = True

                print(y[0])
                index = index + 1
                # print(ind)

            #else:
            #    if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
            #        ind.append(index)
            #        tt = t
            #        y = re.findall('\s\[.+\]\s', t)
            #        if len(y) != 0:
            #            green = paragraph.add_run(y[0])
            #            paragraph.add_run("\n\n")
            #            green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            #            green.bold = True
            #            # print(y[0])
            #    index = index + 1
            # print(ind)


runner2 = paragraph.add_run("\n\nParent tag/tags\n\n")
runner2.bold = True                              #make it bold
GetParentTags()

runner2 = paragraph.add_run("\n\nChild tag/tags\n\n")
runner2.bold = True
GetChildTags()

#runner2 = paragraph.add_run("\n\nOrphanChild tag/tags\n")
#runner2.bold = True
#GetOrphanTags()

report3.save('report3.docx')
GetOrphanTags()

window = Tk()
window.title('Targest')
window.iconbitmap("/Users/") # mac path
# window.iconbitmap("C:/Users/") # windows path

#columns = ('Parent Tag', 'Info', 'Child Tag')
strWindow = StringVar()
pathLabel = Label(window, textvariable=strWindow, fg='blue')
# pathLabel.grid(row=3, column=1)
# strWindow.set("Directory Path")
treeView = ttk.Treeview(window, selectmode='browse')
# treeView.grid(row=3, column=1, columnspan=4, padx=20, pady=20)
treeView['show'] = 'tree'

def saveFile():
    file = filedialog.asksaveasfilename(initialdir="/Users/adrian/",
                                        filetypes=[("Text files", '*.txt'), ("Word document", ".docx"), ("CSV files", '.csv'), ("All file types", "*.*")],
                                        defaultextension='.txt', title="Save file")
    fileObject = open(file, 'w')
    #      if file is None:
    #         return
    # #     # name = file.name
    # #     # baseName = os.path.basename(name)
    # #     # path = os.path.dirname(name)
    # #     # print(path)
    # #     # print(baseName)
    # #     # doc.save(path + "/" + baseName)
    #      fileText = str(text.get(1.0,END))
    # #     #fileText = input()
    fileObject.write("hi")
    #file.write(fileText)
    fileObject.close()

def loadFile():
    file = filedialog.askopenfilename(initialdir="/Users/adrian/", # for this to work just change to your directory path
                                      filetypes=[
                                          ("Word document", "*.docx"),
                                          ("CSV files", '.csv'),
                                          ("Text files", '*.txt'),
                                          ("All file types", "*.*")
                                      ])
    if(file):
        strWindow.set(file)
        file = open(file, 'r') # this can work for string objects if you change "askopenfile" to "askopenfilename"
        #print(file.read()) # change from 'file' to 'fileObject' for "askopenfilename" implementation method to work
        i = 0
        for data in file:
            treeView.insert("", 'end', iid=i, text=data)
            i = i + 1

# def doNothing():
#     fileWin = Toplevel(window)
#     button = Button(fileWin, text="Do nothing button")
#     button.pack()

# def newFile():
#     myText.delete("1.0", END)
#     window.title('Targest')
#     statusBar.config(text="New File   ")
#

# textScroll.config(command=myText.yview)
text = Text(window, borderwidth=10, background='light grey')
text.pack()
#
# menuBar = Menu(window)
# fileMenu = Menu(menuBar, tearoff=0)
# fileMenu.add_command(label="New", command=doNothing)
# fileMenu.add_command(label="Open", command=doNothing)
# #fileMenu.add_command(label="Save", command=saveFile)
# fileMenu.add_command(label="Save as...", command=doNothing)
# fileMenu.add_command(label="Close", command=doNothing)
#
# fileMenu.add_separator()
#
# fileMenu.add_command(label="Exit", command=window.quit)
# menuBar.add_cascade(label="File", menu=fileMenu)
# editMenu = Menu(menuBar, tearoff=0)
# editMenu.add_command(label="Undo", command=doNothing)
#
# editMenu.add_separator()
#
# editMenu.add_command(label="Cut", command=doNothing)
# editMenu.add_command(label="Copy", command=doNothing)
# editMenu.add_command(label="Paste", command=doNothing)
# editMenu.add_command(label="Delete", command=doNothing)
# editMenu.add_command(label="Select All", command=doNothing)
#
# menuBar.add_cascade(label="Edit", menu=editMenu)
# helpMenu = Menu(menuBar, tearoff=0)
# helpMenu.add_command(label="Help Index", command=doNothing)
# helpMenu.add_command(label="About...", command=doNothing)
# menuBar.add_cascade(label="Help", menu=helpMenu)
#
# window.config(menu=menuBar)
saveButton = Button(window, text='Save', fg='blue', activeforeground='red', width=20, command=saveFile)
loadButton = Button(window, text='Load', fg='blue', activeforeground='red', width=20, command=loadFile)
saveButton.pack(side=LEFT, padx=15, pady=20)
loadButton.pack(side=BOTTOM, padx=15, pady=20)
window.mainloop()

